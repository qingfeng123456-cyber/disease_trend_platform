from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "3.3接口设计.docx"

BLUE = "2E5F85"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "000000"
TABLE_WIDTH_DXA = 8900
TABLE_INDENT_DXA = 120


def set_run_font(run, *, east_asia="宋体", latin="Calibri", size=10.5, bold=False, color=BLACK):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.0):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, *, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for idx, title in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT_BLUE)
        set_cell_text(
            header.cells[idx],
            title,
            bold=True,
            color=DARK_BLUE,
            align=(alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.CENTER),
            size=9.2,
        )
    set_repeat_table_header(header)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(
                row.cells[idx],
                value,
                align=(alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT),
            )
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, *, bold_lead=None):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.1
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="等线", latin="Consolas", size=8.8, color="344054")
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_key_value_table(doc, rows):
    return add_table(
        doc,
        ["接口参数", "值"],
        rows,
        [1800, 7100],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    heading_tokens = {
        1: (16, BLUE, 14, 8),
        2: (13, BLUE, 12, 6),
        3: (11.5, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("系统架构与详细设计说明书  |  接口设计")
    set_run_font(run, east_asia="等线", size=8.5, color=MID_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    field_run = footer.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    display_text = OxmlElement("w:t")
    display_text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char_begin)
    field_run._r.append(instr_text)
    field_run._r.append(fld_char_separate)
    field_run._r.append(display_text)
    field_run._r.append(fld_char_end)
    run = footer.add_run(" 页")
    set_run_font(run, size=8.5, color=MID_GRAY)


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("3.3 接口设计")
    set_run_font(title_run, east_asia="黑体", size=20, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("基于机器学习的传染病发病趋势预测与可视化平台")
    set_run_font(subtitle_run, east_asia="等线", size=11, color=MID_GRAY)

    add_body(
        doc,
        "本系统采用前后端分离的数据交互方式。前端页面通过 Fetch API 向 Flask 后端发送 HTTP 请求，Flask 从 Serving 层 JSON 文件读取数据，经参数校验、疾病和地区筛选、日期范围过滤及统计汇总后，以 JSON 格式返回。数据处理模块之间采用 CSV、JSON 和模型文件作为接口，保证采集、清洗、特征工程、模型训练和 Web 展示能够独立运行并协同工作。",
    )
    add_body(
        doc,
        "本项目定位为公开数据分析与课程教学平台，当前未实现用户登录、注册和权限管理功能，因此不设置登录接口。Web API 主要提供只读查询，不允许浏览器直接修改 Raw、Silver、Gold、Models 或 Serving 层文件。",
    )

    add_heading(doc, "3.3.1 接口总体架构", 2)
    add_body(
        doc,
        "接口体系由外部数据源接口、流水线文件接口、Flask REST API 和前端调用接口四部分组成。外部采集器从 OWID、WHO GHO、Open-Meteo、World Bank 等公开来源获得数据；本地 Pandas 或远程 Spark 作业将数据依次写入 Raw、Silver 和 Gold 层；模型模块读取 Gold 特征并输出模型文件、预测结果和评价指标；Serving 导出程序将页面所需内容转换为 JSON；Flask 的 DataService 读取这些 JSON，并通过 /api/ 路径向 ECharts 页面提供查询服务。",
    )
    add_code_block(
        doc,
        "公开数据源 -> Raw -> Silver -> Gold -> 模型训练与评价\n"
        "           -> Serving JSON -> Flask REST API -> ECharts 浏览器页面",
    )

    add_heading(doc, "3.3.2 通信协议与数据规范", 2)
    add_body(
        doc,
        "系统 Web 接口基于 HTTP 协议，统一采用 GET 方法获取数据，接口路径以 /api/ 开头。请求和响应均使用 UTF-8 编码，响应类型为 application/json。日期参数使用 YYYY-MM-DD 格式；地区参数优先使用 ISO3 国家代码；疾病名称和模型标识必须来自 /api/options 返回的可用选项。",
    )
    add_table(
        doc,
        ["参数", "类型", "是否必填", "说明"],
        [
            ("location", "String", "否", "ISO3 地区代码，如 CHN、USA；接口同时兼容 iso 参数"),
            ("disease", "String", "否", "疾病名称，如 COVID-19、Influenza、Tuberculosis"),
            ("start_date", "String", "否", "查询开始日期，格式为 YYYY-MM-DD"),
            ("end_date", "String", "否", "查询结束日期，格式为 YYYY-MM-DD"),
            ("model", "String", "否", "模型标识，如 naive_last_value、moving_average、local_pytorch_lstm"),
        ],
        [1700, 1100, 1200, 4900],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(doc, "接口成功时使用统一响应信封，业务数据位于 data 字段中：")
    add_code_block(
        doc,
        '{\n  "ok": true,\n  "status": "ok",\n  "data": {},\n  "error": null\n}',
    )
    add_body(doc, "接口失败时返回错误编码和可读错误信息：")
    add_code_block(
        doc,
        '{\n  "ok": false,\n  "status": "error",\n  "data": null,\n  "error": {\n    "code": "validation_error",\n    "message": "start_date 不能晚于 end_date"\n  }\n}',
    )
    add_table(
        doc,
        ["HTTP 状态码", "错误编码", "含义", "处理方式"],
        [
            ("200", "-", "请求成功", "前端读取 data 并刷新页面组件"),
            ("400", "validation_error", "参数格式错误或请求组合无效", "前端显示接口返回的参数提示"),
            ("503", "data_not_ready", "缺少必要 Serving 文件", "先运行真实数据流水线生成 JSON"),
            ("500", "platform_error / internal_error", "平台异常或未处理异常", "记录 logs/platform.log 并返回通用错误提示"),
        ],
        [1300, 2100, 2600, 2900],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "3.3.3 Web API 接口清单", 2)
    add_body(
        doc,
        "Flask 在 src/web/app.py 中定义 14 个只读 API，由 src/web/services/data_service.py 完成 Serving 文件读取、缓存、筛选和汇总。接口清单如下。",
    )
    api_rows = [
        ("系统健康检查", "/api/health", "GET", "无", "状态、数据模式、更新时间和版本"),
        ("综合指标", "/api/overview", "GET", "地区、疾病、日期", "最新值、累计值、风险地区数和最佳模型"),
        ("趋势查询", "/api/trend", "GET", "地区、疾病、日期、模型", "观测、移动平均、预测和参考范围"),
        ("风险地图", "/api/risk-map", "GET", "疾病", "地区坐标、风险分值和风险等级"),
        ("地区排行", "/api/rankings", "GET", "疾病", "风险、增长率和预测排行"),
        ("模型指标", "/api/model-metrics", "GET", "疾病", "MAE、RMSE、R²、MAPE、sMAPE"),
        ("数据质量", "/api/data-quality", "GET", "无", "完整率、缺失率、修订数和匹配统计"),
        ("页面选项", "/api/options", "GET", "无", "疾病、地区、模型、频率和日期范围"),
        ("模型预测", "/api/predictions", "GET", "地区、疾病、日期、模型", "目标日期、真实值、预测值和误差"),
        ("天气关联", "/api/weather-correlation", "GET", "地区、疾病、日期", "温度、湿度、样本量和相关系数"),
        ("疾病占比", "/api/disease-share", "GET", "无", "各疾病清洗后记录数量及占比"),
        ("数据源状态", "/api/source-status", "GET", "无", "来源状态、原始行数、清洗行数和说明"),
        ("WHO 指标", "/api/who-indicators", "GET", "无", "WHO 指标分类、记录数和用途"),
        ("模型覆盖", "/api/model-coverage", "GET", "无", "疾病频率、模型、样本量和特征覆盖"),
    ]
    add_table(
        doc,
        ["接口名称", "URL", "方法", "请求参数", "主要返回内容"],
        api_rows,
        [1450, 1900, 650, 2150, 2750],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "3.3.4 核心接口详细设计", 2)

    add_heading(doc, "3.3.4.1 页面选项接口", 3)
    add_body(doc, "该接口在页面初始化时首先调用，用于取得疾病、地区、模型和日期范围，避免前端写死可选值。")
    add_key_value_table(
        doc,
        [
            ("URL", "/api/options"),
            ("请求方法", "GET"),
            ("请求参数", "无"),
            ("返回数据", "locations、diseases、models、availability、default_model 等"),
            ("前端用途", "初始化筛选控件，并确定每种疾病可用模型和默认日期范围"),
        ],
    )

    add_heading(doc, "3.3.4.2 综合指标接口", 3)
    add_body(doc, "该接口为页面顶部指标卡提供所选疾病和地区的概览信息。")
    add_key_value_table(
        doc,
        [
            ("URL", "/api/overview"),
            ("请求方法", "GET"),
            ("请求参数", "location、disease、start_date、end_date"),
            ("返回数据", "latest_date、current_total_cases、current_total_deaths、current_new_cases、current_rolling_value、high_risk_regions、best_model、best_model_mae"),
            ("处理过程", "标准化地区与疾病 -> 过滤时间范围 -> 获取最后观测点 -> 合并风险和模型摘要"),
        ],
    )

    add_heading(doc, "3.3.4.3 趋势查询接口", 3)
    add_body(doc, "该接口是主趋势图的核心接口，根据疾病原生频率返回历史观测、移动平均和指定模型预测。")
    add_key_value_table(
        doc,
        [
            ("URL", "/api/trend"),
            ("请求方法", "GET"),
            ("请求头", "Accept: application/json"),
            ("请求参数", "location、disease、start_date、end_date、model"),
            ("返回数据", "location_code、disease、frequency、metric_label、model、points、reporting_profile"),
            ("points 结构", "date、actual、rolling_7、prediction、lower、upper 等字段"),
        ],
    )
    add_body(doc, "请求示例：")
    add_code_block(
        doc,
        "/api/trend?location=CHN&disease=COVID-19&start_date=2020-01-22&end_date=2021-05-29&model=local_pytorch_lstm",
    )
    add_body(
        doc,
        "后端首先校验日期顺序，再根据地区代码和疾病名称查找趋势序列，最后根据模型标识将对应预测列转换为统一 prediction 字段。前端把 actual、rolling_7 和 prediction 分别绘制为观测线、移动平均线和预测线。",
    )

    add_heading(doc, "3.3.4.4 模型预测接口", 3)
    add_body(doc, "该接口用于预测误差图和未来目标期预测结果展示。")
    add_key_value_table(
        doc,
        [
            ("URL", "/api/predictions"),
            ("请求方法", "GET"),
            ("请求参数", "location、disease、start_date、end_date、model"),
            ("返回数据", "模型名称及预测记录数组 items"),
            ("预测记录", "观测日期、目标日期、实际值、预测值和 error"),
            ("误差定义", "error = prediction - actual；正值表示高估，负值表示低估"),
        ],
    )

    add_heading(doc, "3.3.4.5 天气关联接口", 3)
    add_body(doc, "该接口提供同期温度、湿度与疾病指标的匹配样本，并在后端计算皮尔逊相关系数。")
    add_key_value_table(
        doc,
        [
            ("URL", "/api/weather-correlation"),
            ("请求方法", "GET"),
            ("请求参数", "location、disease、start_date、end_date"),
            ("返回数据", "items、sample_size、temperature_correlation、humidity_correlation、matched_date_range、fallback_used、message"),
            ("items 结构", "date、temperature_mean、relative_humidity_mean、new_cases_smoothed、metric_label"),
            ("特殊处理", "所选日期无同期天气时，可回退到该疾病和地区全部可匹配天气期，并将 fallback_used 设置为 true"),
        ],
    )

    add_heading(doc, "3.3.4.6 模型评价接口", 3)
    add_body(doc, "该接口按疾病返回所有可用模型在同一时间测试集上的评价结果。")
    add_key_value_table(
        doc,
        [
            ("URL", "/api/model-metrics"),
            ("请求方法", "GET"),
            ("请求参数", "disease"),
            ("返回数据", "best_model、mae、models、comparison 等"),
            ("评价指标", "MAE、RMSE、R²、MAPE 和 sMAPE"),
            ("前端用途", "生成五指标热力图，并在同一疾病内比较不同模型"),
        ],
    )

    add_heading(doc, "3.3.5 模块间文件接口", 2)
    add_body(
        doc,
        "数据流水线模块之间不通过浏览器 API 传递大规模数据，而是使用分层目录中的文件作为稳定接口。各层只读取上游约定的文件，不直接依赖上游模块内部实现。",
    )
    add_table(
        doc,
        ["上游模块", "接口文件或目录", "下游模块", "主要格式", "接口作用"],
        [
            ("数据采集", "data/raw/", "数据清洗", "CSV、JSON", "保留来源原始数据和采集元数据"),
            ("数据清洗", "data/silver/local/", "特征工程", "规范化 CSV", "统一日期、ISO3、疾病、频率和指标口径"),
            ("特征工程", "data/gold/local/forecast_features.csv", "模型训练", "特征 CSV", "提供滞后、移动平均、天气、人口和预测目标"),
            ("模型训练", "data/models/local/", "预测流程", "PT、Joblib", "保存各疾病 LSTM 和 GBDT 参数"),
            ("模型训练", "data/gold/local/*_predictions.csv", "Serving 导出", "预测 CSV", "保存测试预测、目标日期和误差"),
            ("Serving 导出", "data/serving/*.json", "Flask", "JSON", "为 Web API 提供轻量查询数据"),
            ("Flask", "/api/*", "ECharts 页面", "HTTP JSON", "向浏览器提供统一数据服务"),
        ],
        [1350, 2450, 1300, 1200, 2600],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(
        doc,
        "Silver 层至少统一 date、location、location_code、disease、frequency、metric_label、value 和 source；Gold 层增加 lag、rolling、growth、temperature、humidity、population 和 target 等特征；Serving 层将其转换为适合浏览器按疾病、地区和日期查询的 JSON 结构。只要这些字段契约保持稳定，就可以独立替换采集器、清洗作业或模型算法。",
    )

    add_heading(doc, "3.3.6 外部数据源接口", 2)
    add_table(
        doc,
        ["外部来源", "接口方式", "输入", "输出位置", "异常处理"],
        [
            ("OWID", "公开 CSV 下载", "数据地址和本地缓存配置", "data/raw/owid/", "超时、重试、状态码和字段检查"),
            ("Open-Meteo", "HTTP API", "经纬度、起止日期、天气变量", "data/raw/open_meteo/", "分国家分年份下载，失败重试"),
            ("World Bank", "HTTP API", "国家代码、指标代码、年份", "data/raw/world_bank/", "分页读取和缺失结果检查"),
            ("WHO GHO", "公开 API 或本地文件", "指标代码和数据文件", "data/raw/who/", "保留维度字段并进行重复审计"),
            ("Kaggle", "人工下载或 Kaggle CLI", "公开数据集压缩包或 CSV", "data/raw/kaggle/", "本地文件画像和字段验证"),
            ("China CDC", "网页采集与人工复核", "页面地址和附件链接", "data/raw/china_cdc/", "状态记录、去重和人工复核标记"),
        ],
        [1450, 1500, 2300, 1850, 1800],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(
        doc,
        "外部接口仅由采集模块调用，网页不会直接访问第三方数据源。这样可以避免网络波动影响页面展示，并保留可重复清洗和模型训练所需的原始数据快照。采集过程中不得在日志、文档或版本库中打印和提交 API 密钥。",
    )

    add_heading(doc, "3.3.7 前后端交互流程", 2)
    add_body(
        doc,
        "用户首次进入页面时，前端先调用 /api/options 初始化疾病、地区、日期和模型下拉框；随后并行请求 overview、trend、risk-map、rankings、model-metrics、data-quality、weather-correlation、disease-share 和 source-status 等接口。用户改变疾病、地区、日期或模型后，前端重新组装查询参数并刷新相关图表。",
    )
    add_body(
        doc,
        "src/web/static/js/api.js 负责构造 URL、过滤空参数、发送 Fetch 请求和解析统一响应；app.js 维护当前筛选状态与页面刷新流程；charts.js 将接口返回的数据转换为 ECharts 的 xAxis、series、dataset、visualMap 和 dataZoom 配置。接口请求失败时，前端显示 error.message，而不会继续使用结构不完整的数据渲染图表。",
    )
    add_code_block(
        doc,
        "浏览器筛选条件 -> api.js 组装查询参数 -> Flask 路由\n"
        "-> DataService 校验与筛选 -> Serving JSON -> 统一响应\n"
        "-> app.js 更新状态 -> charts.js 生成配置 -> ECharts 重绘",
    )

    add_heading(doc, "3.3.8 接口协同与性能设计", 2)
    add_body(
        doc,
        "Flask 请求期间只读取 Serving 层，不启动 Spark 作业，也不执行模型训练。DataService 根据文件修改时间和可配置缓存周期缓存 JSON，默认缓存时间为 30 秒；当流水线重新生成文件后，缓存会在下一次读取时自动更新。该设计将耗时的数据处理与轻量的页面查询分离，保证接口响应稳定。",
    )
    add_body(
        doc,
        "接口对地区、疾病、模型和日期进行标准化校验，防止前端传入不存在的选项。趋势和预测记录始终按日期排序；不同疾病保留日频、周频或年频原始统计口径；相关系数和模型误差只在数据匹配有效时返回。由此保证数据采集、清洗、训练、服务和可视化模块能够按照明确的数据契约协同运行。",
    )

    props = doc.core_properties
    props.title = "3.3 接口设计"
    props.subject = "基于机器学习的传染病发病趋势预测与可视化平台"
    props.author = "项目组"
    props.keywords = "Flask, REST API, ECharts, 接口设计, 数据流水线"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
