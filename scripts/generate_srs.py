# -*- coding: utf-8 -*-
"""
生成《基于机器学习的传染病趋势预测与可视化分析平台 软件需求规约》.docx
参考模板：1.软件需求规约.doc 的章节结构和格式风格
"""
import os
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement

# ===================================================================
# 工具函数
# ===================================================================

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_paragraph_with_font(doc, text, font_name='宋体', font_size=12, bold=False,
                             alignment=None, color=None, space_before=0, space_after=0,
                             first_line_indent=None):
    """添加带字体格式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p, space_before, space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    return p


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def set_table_cell(cell, text, font_name='宋体', font_size=10.5, bold=False, alignment=None):
    """设置表格单元格文本和格式"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def set_table_style(table):
    """设置表格基本样式 - 全边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tblPr.append(borders)


# ===================================================================
# 主生成函数
# ===================================================================

def generate():
    doc = Document()

    # ---------- 页面设置 ----------
    section = doc.sections[0]
    section.page_width = Cm(21)    # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ---------- 设置默认字体 ----------
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    today_str = date.today().strftime('%Y-%m-%d')

    # ==================== 封面页 ====================
    # 空行
    for _ in range(4):
        add_paragraph_with_font(doc, '', font_size=12, space_after=0)

    # 卷号等
    add_paragraph_with_font(doc, '卷    号', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '卷内编号', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '密    级', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '项目编号: ____________________', font_size=12, space_after=6)

    add_paragraph_with_font(doc, '', font_size=12, space_after=6)

    # 项目名称标题
    add_paragraph_with_font(doc, '基于机器学习的传染病趋势预测与可视化分析平台',
                            font_name='黑体', font_size=22, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=10)

    add_paragraph_with_font(doc, '分类: ____________________', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '使用者: 课程实训学生、项目答辩评委', font_size=12, space_after=6)

    add_paragraph_with_font(doc, '', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '软件需求规约', font_name='黑体', font_size=26, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)

    add_paragraph_with_font(doc, f'Version: 1.0', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_paragraph_with_font(doc, '', font_size=12, space_after=30)
    add_paragraph_with_font(doc, '项 目 承 担 部 门：  ____________________', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '撰  写  人（签名）：  ____________________', font_size=12, space_after=6)
    add_paragraph_with_font(doc, f'完   成   日   期：    {today_str}', font_size=12, space_after=6)

    add_paragraph_with_font(doc, '', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '本文档 使 用部门： ■主管领导     ■项目组', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '                      ■客户（市场）  ■维护人员  ■用户', font_size=12, space_after=6)

    add_paragraph_with_font(doc, '', font_size=12, space_after=12)
    add_paragraph_with_font(doc, '评审负责人（签名）：  ____________________', font_size=12, space_after=6)
    add_paragraph_with_font(doc, '评    审   日  期：      ____________________', font_size=12, space_after=6)

    # 分页 - 文档信息页
    doc.add_page_break()

    # ==================== 文档信息 ====================
    add_paragraph_with_font(doc, '文档信息', font_name='黑体', font_size=16, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    info_data = [
        ('标题:', '基于机器学习的传染病趋势预测与可视化分析平台软件需求规约'),
        ('作者:', '____________________'),
        ('创建日期:', today_str),
        ('上次更新日期:', today_str),
        ('版本:', '1.0'),
        ('部门名称:', '____________________'),
    ]
    for label, value in info_data:
        p = doc.add_paragraph()
        run1 = p.add_run(label)
        run1.font.name = '宋体'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run1.font.size = Pt(12)
        run1.bold = True
        run2 = p.add_run(f'  {value}')
        run2.font.name = '宋体'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run2.font.size = Pt(12)
        set_paragraph_spacing(p, 4, 4)

    doc.add_page_break()

    # ==================== 修订文档历史记录 ====================
    add_heading_styled(doc, '修订文档历史记录', level=1)

    rev_table = doc.add_table(rows=2, cols=4)
    set_table_style(rev_table)
    rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['日期', '版本', '说明', '作者']
    for i, h in enumerate(headers):
        set_table_cell(rev_table.rows[0].cells[i], h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # 灰色背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        rev_table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading)

    rev_history = [
        (today_str, '1.0.' + today_str.replace('-', ''), '正式发布', '____________________'),
    ]
    for row_data in rev_history:
        row = rev_table.add_row()
        for i, val in enumerate(row_data):
            set_table_cell(row.cells[i], val, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ==================== 目录 ====================
    add_heading_styled(doc, '目  录', level=1)
    p = doc.add_paragraph()
    run = p.add_run('（请在 Microsoft Word 中右键此处 → 更新域 → 更新整个目录，即可自动生成目录）')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # TOC 域代码
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run('（打开文档后，右键此处 → 更新域 → 更新整个目录）')
    run4.font.name = '宋体'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run4.font.size = Pt(10.5)

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

    doc.add_page_break()

    # ================================================================
    # 1. 引言
    # ================================================================
    add_heading_styled(doc, '1. 引言', level=1)

    # 1.1 目的
    add_heading_styled(doc, '1.1 目的', level=2)
    add_paragraph_with_font(doc,
        '本文档旨在定义"基于机器学习的传染病趋势预测与可视化分析平台"的软件需求，'
        '作为用户和软件开发人员之间相互了解的基础。具体目的包括：',
        first_line_indent=24)
    purposes = [
        '定义平台总体功能和非功能需求，明确系统边界；',
        '提供系统架构设计和模块划分的依据，作为软件开发人员进行软件结构设计和编码的基础；',
        '作为软件总体测试和项目验收的依据。',
    ]
    for item in purposes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    # 1.2 范围
    add_heading_styled(doc, '1.2 范围', level=2)
    add_paragraph_with_font(doc,
        '本文档适用于"基于机器学习的传染病趋势预测与可视化分析平台"项目，涵盖从公开数据采集、'
        '分布式存储与清洗、特征工程、机器学习预测建模、RESTful API 服务到 ECharts 可视化大屏的'
        '完整功能需求描述。项目面向课程实训教学和答辩演示场景，支持 demo（固定种子演示数据）和 '
        'real（真实公开数据）两种运行模式。', first_line_indent=24)

    # 1.3 定义、首字母缩写词和缩略语
    add_heading_styled(doc, '1.3 定义、首字母缩写词和缩略语', level=2)

    glossary = [
        ('OWID', 'Our World in Data，提供 COVID-19 及其他全球健康数据的开源数据平台'),
        ('HDFS', 'Hadoop Distributed File System，Hadoop 分布式文件系统'),
        ('PySpark', 'Apache Spark 的 Python API，用于大规模数据处理'),
        ('ETL', 'Extract, Transform, Load，数据提取、转换和加载流程'),
        ('GBDT', 'Gradient Boosted Decision Trees，梯度提升决策树（使用 scikit-learn HistGradientBoosting 实现）'),
        ('LSTM', 'Long Short-Term Memory，长短期记忆网络（使用 PyTorch 实现）'),
        ('ARIMA', 'AutoRegressive Integrated Moving Average，自回归积分滑动平均模型'),
        ('ECharts', '百度开源的 JavaScript 数据可视化图表库'),
        ('Flask', 'Python 轻量级 Web 应用框架'),
        ('API', 'Application Programming Interface，应用程序编程接口'),
        ('JSON', 'JavaScript Object Notation，轻量级数据交换格式'),
        ('MAE', 'Mean Absolute Error，平均绝对误差'),
        ('RMSE', 'Root Mean Square Error，均方根误差'),
        ('R²', 'R-squared，决定系数'),
        ('MAPE', 'Mean Absolute Percentage Error，平均绝对百分比误差'),
        ('SMAPE', 'Symmetric Mean Absolute Percentage Error，对称平均绝对百分比误差'),
        ('WHO', 'World Health Organization，世界卫生组织'),
        ('GHO', 'Global Health Observatory，全球卫生观察站'),
        ('ISO3', 'ISO 3166-1 alpha-3，三字母国家/地区代码'),
        ('Parquet', 'Apache Parquet，列式存储文件格式'),
        ('SSH', 'Secure Shell，安全外壳协议'),
        ('SFTP', 'SSH File Transfer Protocol，SSH 文件传输协议'),
        ('RSV', 'Respiratory Syncytial Virus，呼吸道合胞病毒'),
        ('HIV', 'Human Immunodeficiency Virus，人类免疫缺陷病毒'),
        ('AIDS', 'Acquired Immunodeficiency Syndrome，获得性免疫缺陷综合征'),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(12.5)

    set_table_cell(table.rows[0].cells[0], '缩略语', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_cell(table.rows[0].cells[1], '全称 / 说明', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for row in table.rows:
        shading0 = OxmlElement('w:shd')
        shading0.set(qn('w:fill'), 'D9E2F3')
        shading0.set(qn('w:val'), 'clear')
        row.cells[0]._tc.get_or_add_tcPr().append(shading0)
        shading1 = OxmlElement('w:shd')
        shading1.set(qn('w:fill'), 'D9E2F3')
        shading1.set(qn('w:val'), 'clear')
        row.cells[1]._tc.get_or_add_tcPr().append(shading1)

    for abbr, desc in glossary:
        row = table.add_row()
        set_table_cell(row.cells[0], abbr, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_cell(row.cells[1], desc)

    # 1.4 参考资料
    add_heading_styled(doc, '1.4 参考资料', level=2)
    refs = [
        '《基于机器学习的传染病趋势预测与可视化分析平台解决方案》',
        '《传染病趋势预测平台 - 系统架构文档》（docs/architecture.md）',
        '《传染病趋势预测平台 - API 接口文档》（docs/api.md）',
        '《传染病趋势预测平台 - 数据字典》（docs/data_dictionary.md）',
        '《传染病趋势预测平台 - 数据来源说明》（docs/data_sources.md）',
        '《实验室报告（一）~（七）》（report/*.md）',
    ]
    for i, ref in enumerate(refs, 1):
        add_paragraph_with_font(doc, f'[{i}] {ref}', first_line_indent=24)

    doc.add_page_break()

    # ================================================================
    # 2. 软件总体概述
    # ================================================================
    add_heading_styled(doc, '2. 软件总体概述', level=1)
    add_paragraph_with_font(doc,
        '本文档主要定义了基于机器学习的传染病趋势预测与可视化分析平台的需求，'
        '由七大功能模块组成：数据采集模块、数据存储与分层处理模块、特征工程模块、'
        '机器学习预测模块、Flask Web API 模块、ECharts 可视化大屏模块、系统运行与管理模块。',
        first_line_indent=24)

    # 2.1 软件标识
    add_heading_styled(doc, '2.1 软件标识', level=2)
    add_paragraph_with_font(doc, '软件全名称：基于机器学习的传染病趋势预测与可视化分析平台', first_line_indent=24)
    add_paragraph_with_font(doc, '软件英文名称：Disease Trend Prediction & Visualization Platform', first_line_indent=24)
    add_paragraph_with_font(doc, '软件缩称：DTP', first_line_indent=24)
    add_paragraph_with_font(doc, '版本号：1.0.0', first_line_indent=24)

    # 2.2 软件描述
    add_heading_styled(doc, '2.2 软件描述', level=2)

    # 2.2.1 系统属性
    add_heading_styled(doc, '2.2.1 系统属性', level=3)
    add_paragraph_with_font(doc,
        '本系统是一个独立开发的 Web 数据平台。系统将公开传染病、天气、人口和社会经济数据'
        '组织成可重复运行的数据工程流程，完成多疾病趋势分析、基线/机器学习/LSTM 预测、'
        'Flask REST API 服务和 ECharts 可视化展示。系统同时支持 Windows 本地 Pandas + sklearn + PyTorch '
        '完整演示和 Ubuntu HDFS + Spark 大数据教学流程两种运行环境，以及 demo 固定种子备用流程。',
        first_line_indent=24)

    # 2.2.2 开发背景
    add_heading_styled(doc, '2.2.2 开发背景', level=3)
    add_paragraph_with_font(doc,
        '传染病数据来源分散、更新频繁，不同来源的数据格式、频率和覆盖范围各不相同。'
        '在课程实训和教学场景中，需要一个端到端的平台将公开数据采集、大数据存储与处理、'
        '机器学习建模和可视化分析串联成完整的可运行链路，帮助学生理解大数据技术在公共卫生'
        '分析领域的工程化应用。本平台即为此课程实训目的而开发。',
        first_line_indent=24)

    # 2.2.3 软件功能
    add_heading_styled(doc, '2.2.3 软件功能', level=3)
    add_paragraph_with_font(doc,
        '平台核心功能包括：', first_line_indent=24)

    func_modules = [
        '多源数据采集：自动从 OWID、Open-Meteo、World Bank、WHO GHO、中国 CDC 等公开数据源采集疫情、天气、人口、经济和社会数据，保存原始文件和元信息；',
        '数据存储与分层处理：基于 HDFS 的四层数据架构（raw/silver/gold/serving），使用 PySpark 完成数据清洗标准化、去重、负数修正和异常标记；',
        '特征工程：构建滞后特征、滚动统计、增长率、气象特征和人口特征，使用 lead(target, 7) 构造预测目标，严格按日期切分防止数据泄漏；',
        '机器学习预测：支持最近值基线、移动平均基线、GBDT、LSTM 和 ARIMA 五类模型，覆盖六种疾病（COVID-19、流感、RSV、结核病、HIV/AIDS、COVID-19 住院），输出 MAE、RMSE、R²、MAPE、SMAPE 等评估指标；',
        'Flask Web API：提供 15 个 RESTful API 端点，统一 JSON 响应格式，支持缓存和参数校验；',
        'ECharts 可视化大屏：展示趋势曲线、风险地图、排行榜、模型对比、数据质量仪表盘等 10 余种图表，支持疾病/地区/模型筛选联动；',
        '系统运行管理：提供 PowerShell/Bash 脚本编排完整流水线，支持远程集群管理（SSH/SFTP），内置单元测试和流水线验证工具。',
    ]
    for item in func_modules:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_paragraph_with_font(doc, '以下为系统总体功能模块表：', first_line_indent=24, space_before=12)

    # 功能模块表
    func_table = doc.add_table(rows=1, cols=3)
    set_table_style(func_table)
    func_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_cell(func_table.rows[0].cells[0], '功能模块', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_cell(func_table.rows[0].cells[1], '功能项', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_cell(func_table.rows[0].cells[2], '描述', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for hcell in func_table.rows[0].cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        hcell._tc.get_or_add_tcPr().append(shading)

    func_data = [
        ('数据采集', '疫情数据采集', '从 OWID 获取 COVID-19 全球日频疫情数据'),
        ('数据采集', '天气数据采集', '从 Open-Meteo 获取十国代表城市历史气象数据'),
        ('数据采集', '人口经济数据采集', '从 World Bank API 获取人口、城市化率、人均 GDP'),
        ('数据采集', 'WHO 健康数据采集', '从 WHO GHO API 获取结核病、HIV/AIDS 等健康指标'),
        ('数据采集', '中国 CDC 数据采集', '归档中国疾控中心公开疫情页面和附件'),
        ('存储与处理', 'HDFS 分层存储', '四层目录架构：raw/silver/gold/serving/checkpoints'),
        ('存储与处理', 'Spark 清洗流水线', '字段标准化、去重、负数修订、缺失处理、异常标记'),
        ('存储与处理', '数据质量报告', '自动生成记录数、缺失率、重复数、异常值等质量指标'),
        ('特征工程', '时序特征构造', '滞后特征(lag_1/3/7/14)、滚动均值/标准差、增长率'),
        ('特征工程', '多源特征关联', '按 location_code + date 关联疫情、天气、人口、经济数据'),
        ('特征工程', '预测目标构造', '使用 lead(target, 7) 构造 t+7 预测目标，防止数据泄漏'),
        ('模型预测', '基线模型', '最近值基线和移动平均基线，提供预测性能参照'),
        ('模型预测', 'GBDT 模型', '基于 scikit-learn HistGradientBoosting，60 轮迭代'),
        ('模型预测', 'LSTM 模型', '基于 PyTorch，6 种疾病独立训练，支持早停和标准后训练评估'),
        ('模型预测', 'ARIMA 模型', '基于 statsmodels，支持季节性参数 SARIMAX'),
        ('模型预测', '模型评估', 'MAE/RMSE/R²/MAPE/SMAPE，模型对比和基线对比'),
        ('Web API', '数据查询接口', '趋势查询、风险地图、排行榜、模型指标、数据质量、疾病占比等 15 个端点'),
        ('Web API', '统一响应格式', '所有 API 返回 {"ok":bool, "status":"ok"/"error", "data":{}, "error":null}'),
        ('可视化大屏', '趋势与预测图表', '实际病例与预测趋势、7 日移动平均、增长率分析'),
        ('可视化大屏', '风险分析图表', '风险地图、高风险地区排行榜'),
        ('可视化大屏', '模型对比图表', '模型指标对比柱状图、预测误差分布'),
        ('可视化大屏', '关联分析图表', '温湿度与病例关联散点图、疾病占比饼图'),
        ('可视化大屏', '数据质量仪表盘', '缺失率、异常值、数据源状态等质量指标可视化'),
        ('系统管理', '运行模式管理', 'demo（固定随机种子演示）和 real（真实公开数据）双模式'),
        ('系统管理', '脚本编排', 'PowerShell/Shell 一键构建、验证和启动全套流水线'),
        ('系统管理', '远程集群管理', 'SSH/SFTP 项目同步、远程环境检查和 HDFS/Spark 控制'),
        ('系统管理', '测试与验证', 'pytest 单元测试、流水线完整性验证脚本'),
    ]
    for mod, item, desc in func_data:
        row = func_table.add_row()
        set_table_cell(row.cells[0], mod, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_cell(row.cells[1], item)
        set_table_cell(row.cells[2], desc)

    # 2.3 用户的特点
    add_heading_styled(doc, '2.3 用户的特点', level=2)
    add_paragraph_with_font(doc,
        '本平台主要面向以下用户群体：', first_line_indent=24)
    users = [
        '课程实训学生：具备基本 Python 编程和数据分析基础，需要理解大数据处理流程和机器学习建模过程，能够运行脚本和启动 Web 服务，查看可视化结果；',
        '项目答辩评委：关注系统架构的完整性、技术栈的合理性、数据工程的规范性以及展示效果的质量，不需要直接操作代码和服务器；',
        '项目维护人员（可选）：需要了解项目目录结构、配置文件和脚本用法，能够在新环境中重新部署和运行项目。',
    ]
    for u in users:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(u)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    # 2.4 限制与约束
    add_heading_styled(doc, '2.4 限制与约束', level=2)
    constraints = [
        '开发环境约束：项目使用 Python 3.10+ 开发，核心依赖包括 Flask 3.x、PySpark、PyTorch、scikit-learn 等，需要 Conda 虚拟环境管理依赖；',
        '运行环境约束：Windows 本地模式可直接运行完整流水线；Ubuntu 远程模式需要配置 Hadoop HDFS 和 Spark 环境，涉及 Java 环境变量配置；',
        '数据真实性声明：演示数据仅用于课程展示，不代表真实疫情；风险指数仅为课程项目分析指标，不等同于公共卫生部门正式风险等级，不能用于医疗决策；',
        '网络依赖：真实采集模式需要访问外部 API（OWID、Open-Meteo、World Bank、WHO），网络不可用时需使用 demo 模式；',
        '性能约束：Flask API 请求期间不启动 Spark（离线生成 JSON 后 Flask 只读），避免实时大规模计算阻塞 API 响应；',
        '安全约束：远程密码和密钥通过 .env 文件管理，不提交至版本控制；中国 CDC 采集器仅归档无需登录的公开页面和附件。',
    ]
    for c in constraints:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(c)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # 3. 具体需求
    # ================================================================
    add_heading_styled(doc, '3. 具体需求', level=1)

    # ---------- 3.1 多源数据采集模块 ----------
    add_heading_styled(doc, '3.1 多源数据采集模块', level=2)
    add_paragraph_with_font(doc,
        '数据采集模块负责从多个公开数据源自动获取传染病、天气、人口、经济和社会数据，'
        '保存原始文件和采集元信息（来源 URL、SHA256 校验值、采集时间等），确保数据可追溯。',
        first_line_indent=24)

    add_heading_styled(doc, '3.1.1 OWID 疫情数据采集', level=3)
    add_paragraph_with_font(doc,
        '从 Our World in Data 下载 COVID-19 全球日频数据（compact CSV 格式），包含 new_cases、'
        'new_deaths、total_cases、total_deaths、new_cases_smoothed、population 等字段，'
        '覆盖全球 200+ 国家/地区的时间序列。自动保存原始 CSV 文件、来源 URL、字段报告和采集时间戳。'
        '采集器实现于 src/collectors/owid_collector.py。', first_line_indent=24)

    add_heading_styled(doc, '3.1.2 Open-Meteo 天气数据采集', level=3)
    add_paragraph_with_font(doc,
        '根据 config/weather_locations.csv 中配置的十国代表城市经纬度，按年份分批请求 '
        'Open-Meteo Archive API 获取历史气象数据，包括温度（最高/最低/均值）、降水量、'
        '相对湿度、风速和气压。采集器会检查已有文件的 .meta.json 旁路元数据，自动跳过已完成的年份，支持增量采集。'
        '采集器实现于 src/collectors/open_meteo_collector.py。', first_line_indent=24)

    add_heading_styled(doc, '3.1.3 World Bank 人口经济数据采集', level=3)
    add_paragraph_with_font(doc,
        '通过 World Bank API 获取各国年度人口（SP.POP.TOTL）、城市化率（SP.URB.TOTL.IN.ZS）和'
        '人均 GDP（NY.GDP.PCAP.CD）指标数据，覆盖配置中指定的十国范围（2020-2025 年）。'
        '采集器实现于 src/collectors/world_bank_collector.py。', first_line_indent=24)

    add_heading_styled(doc, '3.1.4 WHO GHO 健康数据采集', level=3)
    add_paragraph_with_font(doc,
        '通过 WHO GHO API 分页采集结核病、HIV/AIDS 等健康指标数据，支持按指标代码指定采集范围。'
        '自动处理分页、指标主题分类和去重，保留 WHO 原始记录 ID、指标名称、地区层级、数值、'
        '估计区间和维度信息（年龄、性别等）。采集器实现于 src/collectors/who_collector.py。', first_line_indent=24)

    add_heading_styled(doc, '3.1.5 中国 CDC 公开数据采集', level=3)
    add_paragraph_with_font(doc,
        '归档中国疾病预防控制中心公开页面的标题、发布时间、附件 URL 和本地归档路径。'
        '仅采集无需登录的公开页面和附件，不绕过任何访问限制。HTML/PDF/Excel 解析与下载分开。'
        '采集器实现于 src/collectors/china_cdc_collector.py。', first_line_indent=24)

    add_heading_styled(doc, '3.1.6 Demo 演示数据生成', level=3)
    add_paragraph_with_font(doc,
        '使用固定随机种子生成完整演示所需的疫情、天气、人口、预测、风险、质量报告和 API JSON 数据，'
        '确保每次运行结果一致，适合在不联网或无 Hadoop/Spark 环境下快速验证 Web 页面功能。'
        '实现于 src/collectors/generate_demo_data.py。', first_line_indent=24)

    # 数据源汇总表
    add_paragraph_with_font(doc, '数据源汇总如下表所示：', first_line_indent=24, space_before=12, space_after=6)

    ds_table = doc.add_table(rows=1, cols=5)
    set_table_style(ds_table)
    ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ds_headers = ['数据源', '数据类型', '更新频率', '覆盖范围', '采集器模块']
    for i, h in enumerate(ds_headers):
        set_table_cell(ds_table.rows[0].cells[i], h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        ds_table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading)

    ds_data = [
        ('OWID', 'COVID-19 日频疫情', '实时更新', '全球 200+ 国家/地区', 'owid_collector.py'),
        ('Open-Meteo', '历史气象数据（温度、降水、湿度、风速、气压）', '按需采集', '十国代表城市', 'open_meteo_collector.py'),
        ('World Bank', '人口、城市化率、人均 GDP', '年度', '十国（2020-2025）', 'world_bank_collector.py'),
        ('WHO GHO', '结核病、HIV/AIDS 健康指标', '按需采集', '全球', 'who_collector.py'),
        ('中国 CDC', '公开疫情页面和附件', '按需采集', '中国', 'china_cdc_collector.py'),
        ('Kaggle 数据集', 'COVID-19、结核病、呼吸道数据、历史天气、人口', '静态数据集', '按数据集覆盖范围', 'generate_demo_data.py'),
    ]
    for row_data in ds_data:
        row = ds_table.add_row()
        for i, val in enumerate(row_data):
            set_table_cell(row.cells[i], val)

    doc.add_page_break()

    # ---------- 3.2 数据存储与分层处理模块 ----------
    add_heading_styled(doc, '3.2 数据存储与分层处理模块', level=2)
    add_paragraph_with_font(doc,
        '数据存储与处理模块基于 HDFS 分布式文件系统和 PySpark 大数据处理框架，实现数据的'
        '分层存储、标准化清洗和质量监控。', first_line_indent=24)

    add_heading_styled(doc, '3.2.1 HDFS 分层存储架构', level=3)
    add_paragraph_with_font(doc,
        '系统采用四层数据架构：', first_line_indent=24)

    layers = [
        'Raw 层（/disease_platform/raw）：存放采集器获取的原始 CSV、JSON、HTML 等文件，仅追加和归档，不覆盖原文件，保留采集元数据；',
        'Silver 层（/disease_platform/silver）：存放经 PySpark 清洗标准化后的 Parquet 文件，完成字段统一、类型转换、去重、负数修正和异常标记；',
        'Gold 层（/disease_platform/gold）：存放多源关联和特征工程后的 Parquet 文件，包含滞后、滚动、气象、人口等特征及预测目标列；',
        'Serving 层（/disease_platform/serving）：存放模型预测结果和可视化所需的小型 JSON 快照文件，供 Flask API 直接读取。',
    ]
    for layer in layers:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(layer)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_heading_styled(doc, '3.2.2 PySpark ETL 清洗流水线', level=3)
    add_paragraph_with_font(doc,
        '清洗流水线包含三个核心 Spark 作业：', first_line_indent=24)

    etl_items = [
        'clean_epidemic.py：清洗疫情日频数据，统一字段名称和类型，标准化 ISO3 国家代码，修订负值病例（截断为 0），标记负数修正标志和异常值；',
        'clean_weather.py：清洗天气数据，将多个城市的小时观测按自然日聚合为日均值（温度、湿度、气压、风速），保留日极值和有效观测数；',
        'clean_population.py：清洗人口经济数据，统一 ISO3 代码，处理缺失年份和格式异常；',
    ]
    for item in etl_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_paragraph_with_font(doc,
        '此外还包括 data_quality_report.py 自动生成 Silver/Gold 层数据质量报告，'
        '输出记录数、日期范围、缺失率、重复数、负数修订、异常值等关键质量指标。',
        first_line_indent=24)

    add_paragraph_with_font(doc,
        '关键设计约束：Flask API 请求期间绝不启动 Spark（Spark 离线生成 JSON，Flask 只读 data/serving），'
        '确保 API 响应时间可控。', first_line_indent=24)

    doc.add_page_break()

    # ---------- 3.3 特征工程模块 ----------
    add_heading_styled(doc, '3.3 特征工程模块', level=2)
    add_paragraph_with_font(doc,
        '特征工程模块负责将清洗后的多源数据关联整合，构造时间序列特征和预测目标，'
        '为机器学习模型提供结构化的训练数据。', first_line_indent=24)

    add_heading_styled(doc, '3.3.1 特征构造', level=3)
    add_paragraph_with_font(doc, '特征工程包含以下特征类型：', first_line_indent=24)

    features = [
        '滞后特征（lag）：lag_1、lag_3、lag_7、lag_14，提供不同时间窗口的历史病例信息；',
        '滚动统计特征（rolling）：rolling_mean_3、rolling_mean_7、rolling_mean_14（7 日/14 日滚动均值），rolling_std_7、rolling_std_14（滚动标准差），平滑短期波动；',
        '增长率特征：growth_rate_1（日增长率）、growth_rate_7（周增长率），反映疫情扩散速度；',
        '时间特征：day_of_week、month、is_weekend，捕获周期性和季节性模式；',
        '气象特征：temperature_mean（日均温）、relative_humidity_mean（日平均相对湿度）、precipitation_sum（日降水量），来自 Open-Meteo 采集的天气数据；',
        '人口经济特征：population（人口）、urban_population_ratio（城市化率）、gdp_per_capita（人均 GDP），来自 World Bank 数据；',
        '衍生特征：cases_per_million（每百万人病例）、deaths_per_million（每百万人死亡），标准化不同国家规模的影响；',
        '预测目标：target_t_plus_7，使用 lead(target_column, 7) 构造，日频表示未来第 7 天，周频表示下一周，年频表示下一年。',
    ]
    for f in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_heading_styled(doc, '3.3.2 防数据泄漏设计', level=3)
    add_paragraph_with_font(doc,
        '时间序列数据严格按日期切分为训练集（前 70%）、验证集（中 15%）和测试集（后 15%），'
        '绝不使用随机打乱方式切分。预测目标 target_t_plus_7 使用 lead() 前瞻函数构造，'
        '本身不作为模型输入特征，确保模型在训练时无法"看到"未来信息。', first_line_indent=24)

    add_heading_styled(doc, '3.3.3 多疾病差异化处理', level=3)
    add_paragraph_with_font(doc,
        '不同疾病采用不同的数据频率和特征处理策略：', first_line_indent=24)
    freq_items = [
        'COVID-19：使用日频数据，预测 new_cases_smoothed（平滑新增病例）；',
        '流感（Influenza）：使用周频原始观测值；',
        'RSV：使用周频原始观测值；',
        '结核病（Tuberculosis）：使用年频原始观测值；',
        'HIV/AIDS：使用年频原始观测值；',
        'COVID-19 住院：使用周频原始观测值。',
    ]
    for item in freq_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    doc.add_page_break()

    # ---------- 3.4 机器学习预测模块 ----------
    add_heading_styled(doc, '3.4 机器学习预测模块', level=2)
    add_paragraph_with_font(doc,
        '机器学习预测模块提供多种预测模型的训练、评估和对比功能，覆盖基线模型、传统机器学习、'
        '深度学习和统计方法。', first_line_indent=24)

    add_heading_styled(doc, '3.4.1 基线模型', level=3)
    add_paragraph_with_font(doc,
        '提供两个简单的基线模型作为预测性能参照：', first_line_indent=24)
    baselines = [
        '最近值基线（Naive Baseline）：使用最后已知观测值作为未来所有时间点的预测值，代表最简单的预测策略；',
        '移动平均基线（Moving Average Baseline）：使用近期观测值的移动平均作为预测值。',
    ]
    for b in baselines:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(b)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_heading_styled(doc, '3.4.2 GBDT 梯度提升决策树', level=3)
    add_paragraph_with_font(doc,
        '使用 scikit-learn 的 HistGradientBoostingRegressor，支持最多 60 轮迭代、最大深度为 5 的'
        '决策树。采用基于直方图的梯度提升算法，适合中等规模表格数据的高效训练，主要用于 COVID-19 的'
        '日频预测。模型参数通过 config/settings.yaml 可配（gbt_max_iter、gbt_max_depth）。',
        first_line_indent=24)

    add_heading_styled(doc, '3.4.3 LSTM 长短期记忆网络', level=3)
    add_paragraph_with_font(doc,
        '使用 PyTorch 实现的 LSTM 模型，为六种疾病分别训练独立模型。COVID-19 使用日频 28 天窗口；'
        '其他疾病使用对应频率（周频/年频）。训练支持以下特性：', first_line_indent=24)
    lstm_features = [
        '标准化：每疾病使用独立 StandardScaler，训练后还原到原始单位计算验证集 MAE；',
        '早停机制：最佳权重由原始单位验证集 MAE 判定，连续 patience 轮不改善即停止（默认 5 轮，可配置）；',
        '可配置参数：LstmEpochs（最大轮次，默认 20）、LstmWindow（历史窗口，默认 28）、LstmBatchSize（默认 128）、LstmHiddenSize（默认 32）；',
        '进度显示：训练日志同时显示变换空间的 val_loss 和原始单位的 val_mae；',
        '权重保存：每个疾病保存独立的 .pt 模型文件到 data/models/local/；',
    ]
    for f in lstm_features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_heading_styled(doc, '3.4.4 ARIMA/SARIMAX 模型', level=3)
    add_paragraph_with_font(doc,
        '基于 statsmodels 库的 ARIMA 和 SARIMAX 实现，支持季节性参数配置，'
        '适用于具有明显周期模式的传染病时间序列分析，作为 GBDT 和 LSTM 的补充方法。'
        '实现于 src/models/arima_baseline.py。', first_line_indent=24)

    add_heading_styled(doc, '3.4.5 模型评估体系', level=3)
    add_paragraph_with_font(doc,
        '所有模型统一使用以下评估指标：', first_line_indent=24)
    metrics_list = [
        'MAE（平均绝对误差）：衡量预测值与实际值的平均绝对偏差；',
        'RMSE（均方根误差）：对大误差更敏感的惩罚性指标；',
        'R²（决定系数）：衡量模型解释的数据变异比例；',
        'MAPE（平均绝对百分比误差）：以百分比形式表达误差；',
        'SMAPE（对称平均绝对百分比误差）：对 MAPE 的对称改进版本。',
    ]
    for m in metrics_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(m)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_paragraph_with_font(doc,
        '模型指标将通过 model_metrics.json 和 model_comparison.json 输出至 serving 层，'
        '供 API 和前端使用。复杂模型差于基线时如实显示，不做人为美化。',
        first_line_indent=24)

    doc.add_page_break()

    # ---------- 3.5 Flask Web API 模块 ----------
    add_heading_styled(doc, '3.5 Flask Web API 模块', level=2)
    add_paragraph_with_font(doc,
        'Flask Web API 模块作为系统的后端服务层，提供 RESTful 接口供前端 ECharts 大屏调用。'
        'API 服务器启动时一次性加载 data/serving 目录下的所有 JSON 文件到内存缓存，'
        '请求期间不启动 Spark 或进行大规模计算，确保毫秒级响应。',
        first_line_indent=24)

    add_heading_styled(doc, '3.5.1 API 端点列表', level=3)
    add_paragraph_with_font(doc, '平台提供以下 15 个 REST API 端点：', first_line_indent=24)

    api_table = doc.add_table(rows=1, cols=4)
    set_table_style(api_table)
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_headers = ['方法', '路径', '参数', '说明']
    for i, h in enumerate(api_headers):
        set_table_cell(api_table.rows[0].cells[i], h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        api_table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading)

    api_data = [
        ('GET', '/api/health', '无', '返回服务状态、数据模式、更新时间和版本信息'),
        ('GET', '/api/overview', '无', '返回地区数量、日期范围、累计病例/死亡、高风险地区、最佳模型、数据完整率'),
        ('GET', '/api/trend', 'location, disease, model, start_date, end_date', '返回指定地区、疾病和模型的趋势数据和时间序列'),
        ('GET', '/api/risk-map', '无', '返回各地区经纬度、风险分、风险等级、增长率和预测病例'),
        ('GET', '/api/rankings', '无', '返回风险、增长率和预测病例三类排名'),
        ('GET', '/api/model-metrics', '无', '返回 MAE/RMSE/R²/MAPE/SMAPE、基线对比和特征列表'),
        ('GET', '/api/data-quality', '无', '返回记录数、缺失率、重复数、异常值、来源覆盖等质量指标'),
        ('GET', '/api/options', '无', '返回可选地区列表、疾病列表、模型列表和日期范围'),
        ('GET', '/api/predictions', 'location, disease, model', '返回指定维度的详细预测值和置信区间'),
        ('GET', '/api/weather-correlation', '无', '返回温湿度与病例的关联分析数据'),
        ('GET', '/api/disease-share', '无', '返回各疾病在总病例中的占比数据'),
        ('GET', '/api/source-status', '无', '返回各数据源的采集状态、最后更新时间和记录数'),
        ('GET', '/api/who-indicators', 'indicator, location', '返回 WHO 健康指标查询数据'),
        ('GET', '/api/model-coverage', '无', '返回各模型在各疾病/地区上的覆盖情况矩阵'),
        ('GET', '/api/model-comparison', '无', '返回多模型性能对比数据'),
    ]
    for row_data in api_data:
        row = api_table.add_row()
        for i, val in enumerate(row_data):
            set_table_cell(row.cells[i], val, font_size=9)

    add_heading_styled(doc, '3.5.2 统一响应格式', level=3)
    add_paragraph_with_font(doc, '所有 API 统一返回以下 JSON 结构：', first_line_indent=24)
    add_paragraph_with_font(doc,
        '成功响应：{"ok": true, "status": "ok", "data": { ... }, "error": null}',
        first_line_indent=24, font_size=10.5)
    add_paragraph_with_font(doc,
        '错误响应：{"ok": false, "status": "error", "data": null, "error": {"code": "validation_error", "message": "参数说明"}}',
        first_line_indent=24, font_size=10.5)

    add_heading_styled(doc, '3.5.3 缓存策略', level=3)
    add_paragraph_with_font(doc,
        'API 数据服务（DataService）在 Flask 启动时一次性加载 serving 目录下所有 JSON 文件到内存，'
        '提供基于时间戳的自动缓存刷新机制（默认 cache_seconds=30 秒），'
        '避免每次请求都进行文件 I/O 操作。文件缺失时返回友好的错误信息而不抛出异常。',
        first_line_indent=24)

    doc.add_page_break()

    # ---------- 3.6 ECharts 可视化大屏模块 ----------
    add_heading_styled(doc, '3.6 ECharts 可视化大屏模块', level=2)
    add_paragraph_with_font(doc,
        'ECharts 可视化大屏模块是系统的前端展示层，使用 ECharts 5（CDN 加载）渲染交互式图表，'
        '提供直观的疫情趋势分析和模型对比界面。页面采用暗色 Dashboard 风格，支持响应式布局。',
        first_line_indent=24)

    add_heading_styled(doc, '3.6.1 图表组件列表', level=3)

    chart_table = doc.add_table(rows=1, cols=3)
    set_table_style(chart_table)
    chart_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    chart_headers = ['图表名称', '类型', '说明']
    for i, h in enumerate(chart_headers):
        set_table_cell(chart_table.rows[0].cells[i], h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        chart_table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading)

    chart_data = [
        ('实际病例与预测趋势', '折线图 + 柱状图', '展示实际病例值和模型预测值的对比时间序列，支持多模型叠加'),
        ('7 日移动平均', '折线图', '展示 7 日滚动均值的平滑趋势，降低日间波动干扰'),
        ('风险地图', '地图（散点/热力）', '在地理位置上标注各地区风险等级和风险分值'),
        ('高风险地区排行榜', '柱状图/条形图', '按风险分排序展示高风险地区'),
        ('温湿度与病例关联', '散点图', '展示气象因素（温度、湿度）与病例数的关联关系'),
        ('模型指标对比', '柱状图', '对比多个模型在不同评估指标（MAE/RMSE/R²）上的性能'),
        ('数据质量仪表盘', '仪表盘 + 进度条', '展示数据缺失率、异常值比例、数据源完整度等质量 KPI'),
        ('增长率分析', '折线图', '展示病例日增长率和周增长率的变化趋势'),
        ('疾病占比', '饼图/环形图', '展示不同疾病在总病例中的占比分布'),
        ('预测误差分布', '散点图/折线图', '展示预测值与实际值的误差分布情况'),
    ]
    for row_data in chart_data:
        row = chart_table.add_row()
        for i, val in enumerate(row_data):
            set_table_cell(row.cells[i], val)

    add_heading_styled(doc, '3.6.2 交互功能', level=3)
    interactions = [
        '筛选器联动：顶部提供疾病选择器、地区选择器和模型选择器，切换任一选项后所有图表自动刷新；',
        '时间范围选择：支持通过日期选择器调整趋势图的时间区间；',
        'Tooltip 悬停信息：所有图表支持鼠标悬停显示详细信息（数值、日期、百分比等）；',
        '图例切换：多模型/多疾病场景支持通过图例开关控制显示；',
        '数据真实性标识：页面明确标注演示数据声明和风险指数分析指标性质。',
    ]
    for item in interactions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_heading_styled(doc, '3.6.3 KPI 指标卡片', level=3)
    add_paragraph_with_font(doc,
        '页面顶部展示核心 KPI 指标卡片，包括：累计确诊病例数、累计死亡数、监测地区数、'
        '数据完整率、最佳模型名称等概览信息，数据来源于 /api/overview 接口。',
        first_line_indent=24)

    doc.add_page_break()

    # ---------- 3.7 系统运行与管理模块 ----------
    add_heading_styled(doc, '3.7 系统运行与管理模块', level=2)

    add_heading_styled(doc, '3.7.1 双模式运行', level=3)
    add_paragraph_with_font(doc,
        '平台支持两种运行模式：', first_line_indent=24)
    modes = [
        'Demo 模式：使用固定随机种子生成演示数据，不需要网络、Hadoop 或 Spark，适合先跑通网页和测试前端功能；',
        'Real 模式：采集公开真实数据，通过本地 Pandas 清洗或多机 HDFS+Spark 流水线处理，生成完整的 silver/gold/serving 数据。',
    ]
    for m in modes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(m)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    add_heading_styled(doc, '3.7.2 脚本编排系统', level=3)
    add_paragraph_with_font(doc,
        '项目提供完整的自动化脚本集，支持一键运行全流程：', first_line_indent=24)

    script_table = doc.add_table(rows=1, cols=3)
    set_table_style(script_table)
    script_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    script_headers = ['脚本', '平台/语言', '功能']
    for i, h in enumerate(script_headers):
        set_table_cell(script_table.rows[0].cells[i], h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        script_table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading)

    scripts_data = [
        ('run_local_real_pipeline.ps1', 'Windows PowerShell', '本地完整流水线：清洗→特征→GBDT→LSTM→Serving→Flask 启动'),
        ('install_lstm_dependencies.ps1', 'Windows PowerShell', '安装 PyTorch CPU 依赖'),
        ('build_local_serving_from_raw.py', 'Python', '本地多源清洗、Gold 特征、模型训练和 Serving 生成的总编排'),
        ('run_all.sh', 'Linux Bash', 'Linux 端一键 demo/real 总编排脚本'),
        ('run_spark_pipeline.sh', 'Linux Bash', 'Spark 清洗、特征、GBT、Serving 全链脚本'),
        ('start_web.sh / stop_web.sh', 'Linux Bash', '启动/停止 Flask Web 服务'),
        ('remote_pipeline.py', 'Python', 'Windows 远程控制 Ubuntu 集群：SSH/SFTP 同步与执行'),
        ('verify_local_pipeline.py', 'Python', '检查本地模型/Serving 契约和输出完整性'),
    ]
    for row_data in scripts_data:
        row = script_table.add_row()
        for i, val in enumerate(row_data):
            set_table_cell(row.cells[i], val, font_size=9)

    add_heading_styled(doc, '3.7.3 远程集群管理', level=3)
    add_paragraph_with_font(doc,
        '通过 SSH/SFTP（基于 Paramiko）实现 Windows 开发机对 Ubuntu 远程集群的控制：'
        '项目文件比较和同步（project_sync.py）、远程环境检查和初始化（remote_environment.py）、'
        'HDFS/Spark 服务状态检查与启停（cluster_manager.py）。'
        '所有远程密码和密钥通过 .env 文件管理，不提交至版本控制。',
        first_line_indent=24)

    add_heading_styled(doc, '3.7.4 测试与验证', level=3)
    add_paragraph_with_font(doc,
        '项目包含完整的测试体系：', first_line_indent=24)
    tests = [
        'pytest 单元测试（tests/ 目录）：覆盖 API 端点、清洗规则、国家映射、疫情聚合、特征工程、模型指标、LSTM、HDFS CLI、远程配置、SSH 等 22 个测试文件；',
        '流水线验证（verify_local_pipeline.py）：自动检查 raw/silver/gold/serving 各层数据完整性、模型文件存在性和 API 响应正确性；',
        '外部网络、真实 SSH 和真实 HDFS 相关测试使用 mock，不依赖实际集群连接。',
    ]
    for t in tests:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(t)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # 4. 性能
    # ================================================================
    add_heading_styled(doc, '4. 性能', level=1)

    add_paragraph_with_font(doc,
        '系统性能要求如下：', first_line_indent=24)

    perf_items = [
        'API 响应时间：所有 API 接口的响应时间不应超过 3 秒（Flask 仅读取预生成的 JSON 文件，无实时计算）；',
        '数据缓存：API 数据服务采用内存缓存机制，默认缓存刷新间隔为 30 秒（可通过 config/settings.yaml 的 web.cache_seconds 配置），减少文件 I/O 开销；',
        'Spark 离线处理：PySpark ETL 作业在本地模式下使用 local[*] 主节点配置，shuffle 分区数为 8，driver 和 executor 内存各 2GB，适用于中等规模数据的处理；',
        '模型训练时间：GBDT 模型在数据集规模适当时训练时间应控制在 5 分钟以内；LSTM 模型按疾病独立训练，单疾病最大 20-40 轮次，应支持早停以减少不必要的训练开销；',
        'Web 前端加载：页面使用 ECharts CDN 加载，首次加载应在 5 秒内完成；图表数据采用并行异步请求方式获取，减少串行等待时间；',
        '并发支持：Flask 内置服务器（debug=False）应支持至少 5 个并发浏览请求不出现阻塞；生产环境建议使用 Gunicorn 或 Waitress 部署。',
    ]
    for item in perf_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # 5. 接口
    # ================================================================
    add_heading_styled(doc, '5. 接口', level=1)

    # 5.1 软件接口
    add_heading_styled(doc, '5.1 软件接口', level=2)

    add_heading_styled(doc, '5.1.1 REST API 接口', level=3)
    add_paragraph_with_font(doc,
        '平台对外提供以下 RESTful API 接口，供前端 ECharts 大屏和外部系统调用。'
        '所有接口均返回统一 JSON 格式。', first_line_indent=24)

    # 详细 API 接口表
    api_detail = doc.add_table(rows=1, cols=5)
    set_table_style(api_detail)
    api_detail.alignment = WD_TABLE_ALIGNMENT.CENTER
    detail_headers = ['编号', '接口路径', '方法', '请求参数', '返回数据说明']
    for i, h in enumerate(detail_headers):
        set_table_cell(api_detail.rows[0].cells[i], h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        api_detail.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading)

    api_details = [
        ('1', '/api/health', 'GET', '无', '服务状态(status)、数据模式(mode)、更新时间(updated_at)、版本(version)'),
        ('2', '/api/overview', 'GET', '无', '地区数、日期范围、累计病例/死亡、高风险地区、数据完整率、免责声明'),
        ('3', '/api/trend', 'GET', 'location(ISO3代码)、disease(疾病名)、model(模型名)、start_date、end_date(可选)', '日期、实际病例、预测值、模型名称的时间序列数组'),
        ('4', '/api/risk-map', 'GET', '无', '各地区经纬度、风险分(risk_score)、风险等级(risk_level)、增长率、预测病例的地图数据'),
        ('5', '/api/rankings', 'GET', '无', '风险排名、增长率排名、预测病例排名三类排行榜'),
        ('6', '/api/model-metrics', 'GET', '无', 'MAE/RMSE/R²/MAPE/SMAPE指标值、基线MAE、特征列表(feature_list)'),
        ('7', '/api/data-quality', 'GET', '无', '记录数、缺失率、重复数、负数修订、异常值、关联失败数、来源覆盖记录数'),
        ('8', '/api/options', 'GET', '无', '可选地区列表、疾病列表、模型列表、日期范围（供前端筛选器使用）'),
        ('9', '/api/predictions', 'GET', 'location, disease, model', '详细预测值、置信区间、预测时间序列'),
        ('10', '/api/weather-correlation', 'GET', '无', '温度、湿度与病例数的关联分析散点数据'),
        ('11', '/api/disease-share', 'GET', '无', '各疾病类型在总病例中的占比（饼图数据）'),
        ('12', '/api/source-status', 'GET', '无', '各采集器状态、最后更新时间、记录数、是否有错误'),
        ('13', '/api/who-indicators', 'GET', 'indicator(指标代码)、location(地区)', 'WHO 指标值、年份、估计区间、维度信息'),
        ('14', '/api/model-coverage', 'GET', '无', '各模型在各疾病/地区上的可用性矩阵'),
        ('15', '/api/model-comparison', 'GET', '无', '多模型的 MAE/RMSE/R² 对比数据和排名'),
    ]
    for row_data in api_details:
        row = api_detail.add_row()
        for i, val in enumerate(row_data):
            set_table_cell(row.cells[i], val, font_size=8)

    add_heading_styled(doc, '5.1.2 统一响应格式', level=3)
    add_paragraph_with_font(doc, '所有 API 接口遵循统一的 JSON 响应格式：', first_line_indent=24)
    add_paragraph_with_font(doc,
        '成功：{"ok": true, "status": "ok", "data": {具体业务数据}, "error": null}',
        first_line_indent=24, font_size=10.5)
    add_paragraph_with_font(doc,
        '失败：{"ok": false, "status": "error", "data": null, "error": {"code": "错误码", "message": "错误描述"}}',
        first_line_indent=24, font_size=10.5)
    add_paragraph_with_font(doc,
        '错误码包括：validation_error（参数校验失败）、not_found（数据不存在）、'
        'internal_error（服务器内部错误）等。', first_line_indent=24)

    add_heading_styled(doc, '5.1.3 数据源接口', level=3)
    add_paragraph_with_font(doc,
        '平台通过以下外部数据接口获取数据：', first_line_indent=24)
    external_apis = [
        'OWID COVID-19 数据接口：https://catalog.ourworldindata.org/garden/covid/latest/compact/compact.csv，HTTP GET 获取 CSV 文件；',
        'Open-Meteo Archive API：https://archive-api.open-meteo.com/v1/archive，HTTP GET，参数包括 latitude、longitude、start_date、end_date、daily 变量列表；',
        'World Bank API：https://api.worldbank.org/v2/country/{iso3}/indicator/{indicator}?format=json，HTTP GET；',
        'WHO GHO API：https://ghoapi.azureedge.net/api/{indicator}，HTTP GET，支持分页参数 $top/$skip；',
        '中国 CDC 公开页面：https://www.chinacdc.cn/jksj/，HTTP GET 页面采集。',
    ]
    for api in external_apis:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(api)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

    # ==================== 保存 ====================
    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        '软件需求规约_传染病趋势预测平台.docx'
    )
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    generate()
